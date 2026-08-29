from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# compact_reference_guide preset with one named override:
# australian_a4_resume = A4 portrait, 0.75-inch margins, no page furniture,
# no tables/text boxes/icons, Arial body and restrained Hoju Compass colours.
NAVY = RGBColor(0x1A, 0x27, 0x44)
GOLD = RGBColor(0x80, 0x65, 0x15)
INK = RGBColor(0x20, 0x26, 0x36)
MUTED = RGBColor(0x5E, 0x66, 0x76)


def set_run_font(run, *, size: float, color: RGBColor = INK, bold: bool = False, italic: bool = False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def set_keep_together(paragraph):
    paragraph.paragraph_format.keep_together = True


def add_custom_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    alignment = OxmlElement("w:lvlJc")
    alignment.set(qn("w:val"), "left")
    level.append(alignment)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(number)
    p_pr.append(num_pr)


def configure_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    title = document.styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = NAVY
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True

    heading = document.styles["Heading 1"]
    heading.font.name = "Arial"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    heading.font.size = Pt(11)
    heading.font.bold = True
    heading.font.color.rgb = NAVY
    heading.paragraph_format.space_before = Pt(9)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.keep_with_next = True

    if "Resume Entry" not in document.styles:
        entry = document.styles.add_style("Resume Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        entry = document.styles["Resume Entry"]
    entry.base_style = normal
    entry.font.name = "Arial"
    entry._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    entry._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    entry.font.size = Pt(10.5)
    entry.font.bold = True
    entry.font.color.rgb = INK
    entry.paragraph_format.space_before = Pt(3)
    entry.paragraph_format.space_after = Pt(1)
    entry.paragraph_format.keep_with_next = True

    if "Resume Meta" not in document.styles:
        meta = document.styles.add_style("Resume Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = document.styles["Resume Meta"]
    meta.base_style = normal
    meta.font.name = "Arial"
    meta._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    meta._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    meta.font.size = Pt(9.5)
    meta.font.color.rgb = MUTED
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.keep_with_next = True


def add_section_heading(document: Document, text: str):
    return document.add_paragraph(text.upper(), style="Heading 1")


def add_body(document: Document, text: str, *, italic: bool = False):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=MUTED if italic else INK, italic=italic)
    set_keep_together(paragraph)
    return paragraph


def add_entry(document: Document, title: str, employer: str, dates: str, bullets: list[str], num_id: int):
    entry = document.add_paragraph(style="Resume Entry")
    entry_run = entry.add_run(f"{title} | {employer}")
    set_run_font(entry_run, size=10.5, color=INK, bold=True)
    meta = document.add_paragraph(style="Resume Meta")
    meta_run = meta.add_run(dates)
    set_run_font(meta_run, size=9.5, color=MUTED)
    for item in bullets:
        paragraph = document.add_paragraph()
        apply_bullet(paragraph, num_id)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.08
        set_keep_together(paragraph)
        run = paragraph.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def build(output_path: Path):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    configure_styles(document)
    bullet_num_id = add_custom_bullet_numbering(document)

    name = document.add_paragraph(style="Title")
    name_run = name.add_run("[YOUR NAME]")
    set_run_font(name_run, size=24, color=NAVY, bold=True)

    role = document.add_paragraph()
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(3)
    role_run = role.add_run("[TARGET ROLE]")
    set_run_font(role_run, size=11, color=GOLD, bold=True)
    set_keep_with_next(role)

    contact = document.add_paragraph(style="Resume Meta")
    contact.paragraph_format.space_after = Pt(6)
    contact_run = contact.add_run("[Phone] | [Professional email] | [City, State] | [LinkedIn or portfolio]")
    set_run_font(contact_run, size=9.5, color=MUTED)

    add_section_heading(document, "Professional Summary")
    add_body(document, "[Write 2-3 sentences: target role, relevant experience, strongest verified skills, and the value you can bring.]")

    add_section_heading(document, "Core Skills")
    add_body(document, "[Skill 1] | [Skill 2] | [Skill 3] | [Skill 4] | [Skill 5] | [Skill 6]")

    add_section_heading(document, "Employment Experience")
    add_entry(
        document,
        "[JOB TITLE]",
        "[EMPLOYER], [CITY, STATE]",
        "[Month Year - Present]",
        [
            "[Start with an action verb, describe what you did, and add a result you can verify.]",
            "[Show a relevant responsibility using the employer's or industry's plain language.]",
            "[Add a number only when you can explain where it came from.]",
        ],
        bullet_num_id,
    )
    add_entry(
        document,
        "[PREVIOUS JOB TITLE]",
        "[EMPLOYER], [CITY, STATE]",
        "[Month Year - Month Year]",
        [
            "[Describe one relevant action and its verified result.]",
            "[Describe one customer, safety, quality, process, or teamwork contribution.]",
        ],
        bullet_num_id,
    )

    add_section_heading(document, "Education & Training")
    education = document.add_paragraph(style="Resume Entry")
    education_run = education.add_run("[QUALIFICATION OR COURSE] | [INSTITUTION]")
    set_run_font(education_run, size=10.5, color=INK, bold=True)
    education_meta = document.add_paragraph(style="Resume Meta")
    education_meta_run = education_meta.add_run("[Year completed or expected] | [City, State or Online]")
    set_run_font(education_meta_run, size=9.5, color=MUTED)

    add_section_heading(document, "Licences & Certifications")
    add_body(document, "[Licence or certificate] | [Issuer] | [Expiry date if relevant]")

    add_section_heading(document, "References")
    add_body(document, "Available upon request.")

    document.core_properties.title = "Free Australian Resume Template"
    document.core_properties.subject = "ATS-friendly editable resume template for Australian job applications"
    document.core_properties.keywords = "Australian resume, resume template, ATS, job application"
    document.core_properties.author = "Hoju Compass"
    document.core_properties.last_modified_by = "Hoju Compass"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.output)


if __name__ == "__main__":
    main()
